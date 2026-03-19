from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, func
from typing import List
import json

from app.core.database import get_db
from app.models.database import Conversation, Message, Document
from app.schemas.chat import (
    ChatRequest, ChatResponse, MessageResponse,
    ConversationResponse, ConversationWithMessages,
    DocumentUpload
)
from app.services.chat_service import chat_service
from app.services.rag_service import rag_service
from app.core.config import settings
import logging

logger = logging.getLogger(__name__)
router = APIRouter()


@router.post("/chat", response_model=ChatResponse)
async def chat(
    request: ChatRequest,
    db: AsyncSession = Depends(get_db)
):
    """Endpoint principal pour le chat"""
    try:
        # Créer ou récupérer la conversation
        if request.conversation_id:
            result = await db.execute(
                select(Conversation).where(Conversation.id == request.conversation_id)
            )
            conversation = result.scalar_one_or_none()
            if not conversation:
                raise HTTPException(status_code=404, detail="Conversation not found")
        else:
            # Créer une nouvelle conversation
            conversation = Conversation(
                title=request.message[:50] + "..." if len(request.message) > 50 else request.message,
                user_id=request.user_id
            )
            db.add(conversation)
            await db.commit()
            await db.refresh(conversation)
        
        # Enregistrer le message utilisateur
        user_message = Message(
            conversation_id=conversation.id,
            role="user",
            content=request.message
        )
        db.add(user_message)
        await db.commit()
        
        # Récupérer l'historique de conversation
        result = await db.execute(
            select(Message)
            .where(Message.conversation_id == conversation.id)
            .order_by(Message.created_at)
        )
        history = result.scalars().all()
        
        # Préparer les messages
        conversation_history = [
            {"role": msg.role, "content": msg.content}
            for msg in history[:-1]  # Exclure le dernier message (déjà dans request.message)
        ]
        
        messages = chat_service.prepare_messages(
            conversation_history=conversation_history,
            new_message=request.message
        )
        
        # Générer la réponse
        if request.stream:
            # Pour le streaming, on retourne un StreamingResponse
            async def generate():
                full_response = ""
                async for chunk in chat_service.generate_response_stream(
                    messages=messages,
                    use_rag=request.use_rag,
                    user_id=request.user_id
                ):
                    full_response += chunk
                    yield f"data: {json.dumps({'chunk': chunk})}\n\n"
                
                # Enregistrer la réponse complète
                assistant_message = Message(
                    conversation_id=conversation.id,
                    role="assistant",
                    content=full_response
                )
                db.add(assistant_message)
                await db.commit()
                
                yield f"data: {json.dumps({'done': True, 'conversation_id': conversation.id})}\n\n"
            
            return StreamingResponse(generate(), media_type="text/event-stream")
        
        else:
            # Génération normale
            response = await chat_service.generate_response(
                messages=messages,
                use_rag=request.use_rag,
                user_id=request.user_id
            )
            
            # Enregistrer la réponse
            assistant_message = Message(
                conversation_id=conversation.id,
                role="assistant",
                content=response
            )
            db.add(assistant_message)
            await db.commit()
            
            return ChatResponse(
                message=response,
                conversation_id=conversation.id,
                sources=None  # TODO: ajouter les sources RAG si pertinent
            )
    
    except Exception as e:
        logger.error(f"Error in chat endpoint: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/conversations", response_model=List[ConversationResponse])
async def get_conversations(
    user_id: str = "anonymous",
    limit: int = 50,
    db: AsyncSession = Depends(get_db)
):
    """Récupère la liste des conversations"""
    result = await db.execute(
        select(Conversation, func.count(Message.id).label('message_count'))
        .outerjoin(Message)
        .where(Conversation.user_id == user_id)
        .group_by(Conversation.id)
        .order_by(Conversation.updated_at.desc())
        .limit(limit)
    )
    
    conversations = []
    for conv, msg_count in result:
        conversations.append(
            ConversationResponse(
                id=conv.id,
                title=conv.title,
                user_id=conv.user_id,
                created_at=conv.created_at,
                updated_at=conv.updated_at,
                message_count=msg_count or 0
            )
        )
    
    return conversations


@router.get("/conversations/{conversation_id}", response_model=ConversationWithMessages)
async def get_conversation(
    conversation_id: str,
    db: AsyncSession = Depends(get_db)
):
    """Récupère une conversation avec tous ses messages"""
    result = await db.execute(
        select(Conversation).where(Conversation.id == conversation_id)
    )
    conversation = result.scalar_one_or_none()
    
    if not conversation:
        raise HTTPException(status_code=404, detail="Conversation not found")
    
    # Récupérer les messages
    result = await db.execute(
        select(Message)
        .where(Message.conversation_id == conversation_id)
        .order_by(Message.created_at)
    )
    messages = result.scalars().all()
    
    return ConversationWithMessages(
        id=conversation.id,
        title=conversation.title,
        user_id=conversation.user_id,
        created_at=conversation.created_at,
        updated_at=conversation.updated_at,
        message_count=len(messages),
        messages=[
            MessageResponse(
                id=msg.id,
                content=msg.content,
                role=msg.role,
                conversation_id=msg.conversation_id,
                created_at=msg.created_at
            )
            for msg in messages
        ]
    )


@router.delete("/conversations/{conversation_id}")
async def delete_conversation(
    conversation_id: str,
    db: AsyncSession = Depends(get_db)
):
    """Supprime une conversation"""
    result = await db.execute(
        select(Conversation).where(Conversation.id == conversation_id)
    )
    conversation = result.scalar_one_or_none()
    
    if not conversation:
        raise HTTPException(status_code=404, detail="Conversation not found")
    
    await db.delete(conversation)
    await db.commit()
    
    return {"status": "success", "message": "Conversation deleted"}


@router.post("/documents/upload")
async def upload_document(
    file: UploadFile = File(...),
    user_id: str = "anonymous",
    db: AsyncSession = Depends(get_db)
):
    """Upload et traite un document pour le RAG"""
    if not rag_service:
        raise HTTPException(status_code=400, detail="RAG is not enabled")
    
    # Vérifier l'extension
    file_ext = "." + file.filename.split(".")[-1].lower()
    if file_ext not in settings.ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"File type not allowed. Allowed: {settings.ALLOWED_EXTENSIONS}"
        )
    
    # Lire le contenu
    content = await file.read()
    
    # Extraire le texte selon le type de fichier
    if file_ext == ".pdf":
        from pypdf import PdfReader
        from io import BytesIO
        pdf = PdfReader(BytesIO(content))
        text = "\n".join(page.extract_text() for page in pdf.pages)
    elif file_ext == ".docx":
        from docx import Document as DocxDocument
        from io import BytesIO
        doc = DocxDocument(BytesIO(content))
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    else:
        text = content.decode("utf-8")
    
    # Créer le document en base
    document = Document(
        filename=file.filename,
        user_id=user_id,
        content=text[:1000]  # Preview
    )
    db.add(document)
    await db.commit()
    await db.refresh(document)
    
    # Ajouter au RAG
    chunk_count = rag_service.add_document(
        document_id=document.id,
        text=text,
        metadata={"user_id": user_id, "filename": file.filename}
    )
    
    # Mettre à jour le compteur de chunks
    document.chunk_count = chunk_count
    await db.commit()
    
    return {
        "status": "success",
        "document_id": document.id,
        "filename": file.filename,
        "chunks": chunk_count
    }


@router.get("/documents")
async def get_documents(
    user_id: str = "anonymous",
    db: AsyncSession = Depends(get_db)
):
    """Récupère la liste des documents uploadés"""
    result = await db.execute(
        select(Document)
        .where(Document.user_id == user_id)
        .order_by(Document.created_at.desc())
    )
    documents = result.scalars().all()
    
    return [
        {
            "id": doc.id,
            "filename": doc.filename,
            "chunk_count": doc.chunk_count,
            "created_at": doc.created_at
        }
        for doc in documents
    ]
